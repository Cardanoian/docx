import os
import google.generativeai as genai
import PyPDF2
import io
import platform
import re
from typing import Tuple

from docx import Document
from docx.shared import Pt

from utils import get_api_key

# --- Initialization ---
API_KEY = get_api_key()

if not API_KEY:
    raise ValueError("GEMINI_API_KEY가 .env 파일에 설정되지 않았습니다.")

genai.configure(api_key=API_KEY)

# --- Configuration ---
GENERATION_CONFIG = {
    "temperature": 0.7,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 10000,
    "response_mime_type": "text/plain",
}

# --- Model Instantiation ---
# Using a recommended model, adjust if needed
llm_model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",  # or "gemini-pro" if preferred
    generation_config=GENERATION_CONFIG,
)


# --- Font Path Setup ---
# 운영체제별 한글 폰트 설정
def get_system_font():
    system = platform.system()
    if system == "Darwin":  # macOS
        return "AppleGothic"
    elif system == "Linux":
        return "NanumGothic"
    elif system == "Windows":
        return "Malgun Gothic"
    else:
        return "Arial"  # 기본 폰트


SYSTEM_FONT = get_system_font()

# model.py 파일이 위치한 디렉토리를 기준으로 fonts 폴더 경로 설정
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FONT_ROOT = os.path.join(BASE_DIR, "fonts")


# --- Core Logic Functions ---
def extract_text_from_pdf(uploaded_pdf_file):
    """PDF 파일 객체에서 텍스트를 추출합니다."""
    text = ""
    if uploaded_pdf_file is None:
        return None
    try:
        # Reset buffer position for reading
        uploaded_pdf_file.seek(0)
        pdf_reader = PyPDF2.PdfReader(uploaded_pdf_file)
        for page in pdf_reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
            except Exception:
                # Log or handle page-specific extraction errors if needed
                continue  # Continue to next page even if one fails
        return text if text else None  # Return None if no text extracted
    except Exception as e:
        print(f"PDF 텍스트 추출 오류: {e}")  # Log error
        return None  # Return None on error


def generate_content_from_gemini(template_text, reference_text, instructions):
    """Gemini 모델을 사용하여 콘텐츠 생성을 스트리밍 방식으로 처리합니다."""
    prompt_text = f"""
# 계획서 또는 보고서 작성

## PDF 서식 파일 내용:
{template_text}

"""
    if reference_text:
        prompt_text += f"""
## 참고 파일 PDF 내용:
{reference_text}

"""

    prompt_text += f"""
## 사용자 지시사항:
{instructions}

---

**지시사항에 따라 PDF 템플릿{", 참고 PDF" if reference_text else ""}를 분석하고, 내용을 채워 계획서 또는 보고서를 작성하세요.**
**한국어로 작성하며, 명확하고 논리적인 구조로 작성해주세요.**
**템플릿 양식에 맞춰 내용을 작성하고, 필요하다면 추가적인 정보나 내용을 생성해도 좋습니다.**
**만약 템플릿 내용이 부족하거나 지시사항을 수행하기 어렵다면, 솔직하게 답변해주세요.**
**학교 사업 계획서, 교육 활동 계획서, 프로젝트 학습 계획서 등 교육 관련 계획서 및 보고서 작성에 특화되어 있습니다.**
**표(테이블)가 필요한 경우 마크다운 테이블 형식으로 생성해주세요.**
"""
    try:
        # stream=True로 설정하여 응답을 청크 단위로 받음
        response_stream = llm_model.generate_content([prompt_text], stream=True)
        for chunk in response_stream:
            # Check if the chunk has text content and it's not empty
            if hasattr(chunk, "text") and chunk.text:
                yield chunk.text  # 각 텍스트 청크를 반환 (yield)
    except Exception as e:
        print(f"Gemini API 호출 오류: {e}")  # Log error
        yield f"\n\n오류 발생: 콘텐츠 생성 중 문제가 발생했습니다. ({e})"  # Yield error message


def markdown_to_docx(markdown_text: str) -> Tuple[str, io.BytesIO]:
    """마크다운 텍스트를 docx 문서로 변환하여 BytesIO로 반환합니다."""
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Malgun Gothic"  # 시스템에 맞게 조정 가능
    font.size = Pt(11)

    lines = markdown_text.strip().splitlines()
    table_mode = False
    table_rows = []
    tables = []
    title = ""

    def parse_inline_styles(paragraph, text):
        """텍스트 내 인라인 스타일 처리: **bold**, *italic*, ***both***"""
        pattern = r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)"
        parts = re.split(pattern, text)
        for part in parts:
            run = paragraph.add_run(re.sub(r"[*]", "", part))  # 기본 텍스트
            if part.startswith("***") and part.endswith("***"):
                run.bold = True
                run.italic = True
            elif part.startswith("**") and part.endswith("**"):
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run.italic = True

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # 제목
        if line.startswith("#"):
            level = min(line.count("#"), 4)
            text = line.lstrip("#").strip()
            doc.add_heading(text, level=level)
            table_mode = False

        # 리스트
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            table_mode = False

        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
            table_mode = False

        # 테이블 감지
        elif "|" in line:
            row = [cell.strip() for cell in line.split("|") if cell.strip()]
            table_rows.append(row)

            # 다음 줄이 헤더 구분줄(---)인지 확인
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", next_line):
                    i += 1  # 구분선은 건너뛰기
                    table_rows.append("__HEADER__")  # 마커로 표시

        else:
            # 이전 테이블 처리
            if table_rows:
                header_style = []
                if "__HEADER__" in table_rows:
                    header_index = table_rows.index("__HEADER__")
                    headers = table_rows[header_index - 1]
                    data_rows = table_rows[header_index + 1 :]
                    all_rows = [headers] + data_rows
                    header_style = [True] + [False] * len(data_rows)
                else:
                    all_rows = table_rows
                    header_style = [False] * len(all_rows)

                num_rows = len(all_rows)
                num_cols = max(len(row) for row in all_rows)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = "Table Grid"

                for r, row in enumerate(all_rows):
                    for c, cell in enumerate(row):
                        cell_text = cell
                        cell_obj = table.cell(r, c)
                        cell_obj.text = cell_text
                        for run in cell_obj.paragraphs[0].runs:
                            run.font.name = "Malgun Gothic"
                            if header_style[r]:
                                run.bold = True
                table_rows = []

            # 일반 문단
            p = doc.add_paragraph()
            parse_inline_styles(p, line)

        if not title:
            title = get_title(line)
        i += 1

    # 혹시 마지막 줄이 테이블이면 처리
    if table_rows:
        all_rows = table_rows
        num_rows = len(all_rows)
        num_cols = max(len(row) for row in all_rows)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        for r, row in enumerate(all_rows):
            for c, cell in enumerate(row):
                cell_obj = table.cell(r, c)
                cell_obj.text = cell
                for run in cell_obj.paragraphs[0].runs:
                    run.font.name = "Malgun Gothic"

    # 결과 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return (title, buffer)


def get_title(line):
    match = re.match(r"^#+\s*(.*)", line)  # Heading tag
    if match:
        return match.group(1).strip()

    match = re.match(r"^-+\s*(.*)", line)  # Unordered list tag
    if match:
        return match.group(1).strip()

    match = re.match(r"^\d+\.\s*(.*)", line)  # Ordered list tag
    if match:
        return match.group(1).strip()

    match = re.match(r"^>\s*(.*)", line)  # Blockquote tag
    if match:
        return match.group(1).strip()

    match = re.match(r"^`{3}.*", line)  # Code block start
    if match:
        return ""  # Code block 시작 줄은 텍스트로 간주하지 않음

    match = re.match(r"^\|.*\|$", line)  # Table row
    if match:
        # 테이블 행 내부의 텍스트를 추출 (간단하게 처리)
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        return " ".join(cells)

    match = re.match(r"^[*_]{1,3}(.*?)[*_]{1,3}$", line)  # Bold, italic
    if match:
        return match.group(1).strip()

    return line.strip()
